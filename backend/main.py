import os
import json
import re
import tempfile
import traceback
import copy
import shutil
from pathlib import Path
from dotenv import load_dotenv
load_dotenv(Path(__file__).parent / ".env", override=True)
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from auth import (authenticate_user, create_token, decode_token, log_generation,
                  add_user, delete_user, get_all_users, get_stats, get_user)
import anthropic
import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

TEMPLATE_PATH = Path(__file__).parent.parent / "template" / "template_infotel.docx"
FRONTEND_PATH = Path(__file__).parent.parent / "frontend" / "index.html"
ADMIN_PATH = Path(__file__).parent.parent / "frontend" / "admin.html"
security = HTTPBearer()


# ─── Auth helpers ─────────────────────────────────────────────────────────────

class LoginRequest(BaseModel):
    username: str
    password: str

class AddUserRequest(BaseModel):
    username: str
    password: str
    role: str = "user"


def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    username = decode_token(credentials.credentials)
    if not username:
        raise HTTPException(401, "Token invalide ou expiré")
    user = get_user(username)
    if not user:
        raise HTTPException(401, "Utilisateur introuvable")
    return user


def require_admin(user=Depends(get_current_user)):
    if user["role"] != "admin":
        raise HTTPException(403, "Accès réservé à l'admin")
    return user


# ─── Routes Auth ──────────────────────────────────────────────────────────────

@app.post("/auth/login")
def login(req: LoginRequest):
    user = authenticate_user(req.username, req.password)
    if not user:
        raise HTTPException(401, "Identifiant ou mot de passe incorrect")
    token = create_token(user["username"])
    return {"token": token, "username": user["username"], "role": user["role"]}


@app.get("/auth/me")
def me(user=Depends(get_current_user)):
    return {"username": user["username"], "role": user["role"],
            "total_generated": user.get("total_generated", 0)}


# ─── Routes Admin ─────────────────────────────────────────────────────────────

@app.get("/admin/stats")
def admin_stats(user=Depends(require_admin)):
    return get_stats()


@app.get("/admin/users")
def admin_users(user=Depends(require_admin)):
    users = get_all_users()
    return [{"username": u["username"], "role": u["role"],
             "total_generated": u.get("total_generated", 0),
             "monthly_generated": u.get("monthly_generated", 0),
             "last_login": u.get("last_login"),
             "created_at": u.get("created_at")} for u in users]


@app.post("/admin/users")
def admin_add_user(req: AddUserRequest, user=Depends(require_admin)):
    ok, msg = add_user(req.username, req.password, req.role)
    if not ok:
        raise HTTPException(400, msg)
    return {"message": "Utilisateur créé"}


@app.delete("/admin/users/{username}")
def admin_delete_user(username: str, user=Depends(require_admin)):
    ok, msg = delete_user(username)
    if not ok:
        raise HTTPException(400, msg)
    return {"message": "Utilisateur supprimé"}
client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))

SYSTEM_PROMPT = """Tu es un expert en analyse de CV pour un cabinet de conseil IT.
Tu réponds UNIQUEMENT avec un objet JSON valide, sans texte avant/après, sans balises markdown.
Règles strictes :
- Ne jamais inventer des informations absentes du CV sauf pour le projet
- Dates au format MM/YYYY ou YYYY
- Missions : extraire TOUTES les missions importantes du CV pour chaque expérience (jusqu'à 12 par expérience). Chaque mission : verbe d'action + résultat concret, 1 ligne max. Ne jamais écrire "Description mission"
- Projets : TOUJOURS renseigner. Si une expérience contient plusieurs projets distincts, les lister tous dans le tableau "projets". Si absent du CV, déduire intelligemment depuis le poste, la société et les missions. Chaque projet : phrase courte et synthétique (max 1 ligne)
- Compétences domaines : exactement 5, phrases courtes et percutantes (max 2 lignes), adaptées au profil
- Compétences techniques : exactement 5 catégories adaptées au profil du consultant
- Expériences significatives : regrouper les postes identiques, lister toutes les sociétés séparées par des virgules, durée totale cumulée (ex: "3 ANS", "18 MOIS")
"""

USER_PROMPT = """Extrais les informations de ce CV et retourne UNIQUEMENT ce JSON (sans markdown) :

{
  "nom": "",
  "prenom": "",
  "annees_experience": "",
  "titre_poste": "",
  "poste": "",
  "competences_domaines": [
    "Phrase 1 max 2 lignes",
    "Phrase 2 max 2 lignes",
    "Phrase 3 max 2 lignes",
    "Phrase 4 max 2 lignes",
    "Phrase 5 max 2 lignes"
  ],
  "competences_techniques": [
    {"categorie": "Catégorie adaptée", "contenu": "Tech1, Tech2, Tech3"},
    {"categorie": "Catégorie adaptée", "contenu": "Tech1, Tech2"},
    {"categorie": "Catégorie adaptée", "contenu": "Tech1, Tech2"},
    {"categorie": "Catégorie adaptée", "contenu": "Tech1, Tech2"},
    {"categorie": "Catégorie adaptée", "contenu": "Tech1, Tech2"}
  ],
  "experiences_significatives": [
    {"titre": "", "societes": "", "duree": ""}
  ],
  "experiences_professionnelles": [
    {
      "date_debut": "",
      "date_fin": "",
      "societe": "",
      "poste": "",
      "projets": ["Projet 1", "Projet 2 si plusieurs"],
      "missions": ["Mission 1", "Mission 2", "Mission 3", "...toutes les missions du CV"],
      "environnement_technique": ""
    }
  ],
  "formations": [
    {"annee": "", "intitule": ""}
  ]
}

CV à analyser :
"""


def extract_text_from_pdf(file_bytes):
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(file_bytes); tmp = f.name
    try:
        with pdfplumber.open(tmp) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)
    finally:
        os.unlink(tmp)


def extract_text_from_docx(file_bytes):
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(file_bytes); tmp = f.name
    try:
        doc = Document(tmp)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    finally:
        os.unlink(tmp)


def build_system_prompt(job_description: str, influence: int) -> str:
    base = SYSTEM_PROMPT
    if not job_description.strip() or influence == 0:
        return base

    influence_instructions = {
        1: (
            "Un besoin client est fourni. Oriente MODÉRÉMENT le CV :\n"
            "- Mets légèrement en avant les expériences et compétences en lien avec ce besoin\n"
            "- Reformule quelques missions pour mieux correspondre au poste\n"
            "- Priorise les compétences techniques pertinentes\n"
            "- Ne jamais inventer de compétences ou expériences absentes du CV"
        ),
        2: (
            "Un besoin client est fourni. Oriente FORTEMENT le CV :\n"
            "- Priorise systématiquement les expériences et compétences en lien avec ce besoin\n"
            "- Reformule les missions pour les aligner au maximum avec le poste visé\n"
            "- Mets en avant uniquement les compétences techniques pertinentes\n"
            "- Adapte le titre du poste si une formulation plus proche du besoin existe dans le CV\n"
            "- Ne jamais inventer de compétences ou expériences absentes du CV"
        ),
    }
    return base + f"\n\nINSTRUCTION D'ORIENTATION :\n{influence_instructions[influence]}"


def extract_cv_data(text: str, job_description: str = "", influence: int = 0):
    system = build_system_prompt(job_description, influence)
    user_content = USER_PROMPT + text
    if job_description.strip() and influence > 0:
        user_content += f"\n\n--- BESOIN CLIENT ---\n{job_description.strip()}"

    print(f"[ORIENTATION] influence={influence}, besoin={'oui ('+str(len(job_description))+' chars)' if job_description.strip() else 'non'}")
    if job_description.strip() and influence > 0:
        print(f"[ORIENTATION] Besoin: {job_description[:100]}...")
    message = client.messages.create(
        model="claude-sonnet-4-5", max_tokens=8000, system=system,
        messages=[{"role": "user", "content": user_content}],
    )
    raw = message.content[0].text.strip()
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


# ─── Helpers XML ──────────────────────────────────────────────────────────────

def get_text_in_element(elem):
    return "".join(t.text or "" for t in elem.iter(qn("w:t")))


def set_text_in_run(run_elem, new_text):
    t_elems = list(run_elem.iter(qn("w:t")))
    if not t_elems: return
    t_elems[0].text = new_text
    t_elems[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    for t in t_elems[1:]: t.text = ""


def _write_para_xml(para, text):
    """Écrit du texte dans le premier run d'un paragraphe via XML direct (fiable sur tout template)."""
    runs = list(para._p.iter(qn("w:r")))
    if not runs:
        return
    set_text_in_run(runs[0], text)
    for r in runs[1:]:
        for t in r.iter(qn("w:t")):
            t.text = ""


def replace_text_in_txbx(doc, marker, replacement):
    body = doc.element.body
    WPS = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx"
    VML = "{urn:schemas-microsoft-com:vml}textbox"
    for elem in body.iter():
        if elem.tag in (WPS, VML):
            for p in elem.iter(qn("w:p")):
                if marker in get_text_in_element(p):
                    runs = list(p.iter(qn("w:r")))
                    if runs:
                        set_text_in_run(runs[0], replacement)
                        for r in runs[1:]: r.getparent().remove(r)


def _is_empty_or_placeholder(text):
    """Détecte les textes vides ou placeholders à effacer."""
    t = text.strip()
    if not t: return True
    placeholders = ["description mission", "compétence", "competence", "techno"]
    return any(p in t.lower() for p in placeholders)


def _remove_bullet_style(para):
    """Supprime la puce et vide le texte d'un paragraphe."""
    for r in para.runs:
        r.text = ""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)


def _clear_bullet_if_empty(para):
    """Efface le contenu d'un paragraphe si vide ou placeholder, en conservant le style."""
    text = "".join(r.text or "" for r in para.runs)
    if _is_empty_or_placeholder(text):
        for r in para.runs:
            r.text = ""
        return True
    return False


def _get_table_header(tbl):
    try: return tbl.rows[0].cells[0].text
    except: return ""


def _is_experience_pro_table(tbl):
    try:
        if len(tbl.columns) == 2:
            c0 = tbl.rows[0].cells[0].text
            c1 = tbl.rows[0].cells[1].text
            if ("20XX" in c0 or "XX 20" in c0) and ("SOCIETE" in c1 or "Poste" in c1 or c1.strip() == ""):
                return True
    except: pass
    return False


def _is_formation_table(tbl):
    try:
        if len(tbl.columns) == 2:
            if "20XX" in tbl.rows[0].cells[0].text and "Formation" in tbl.rows[0].cells[1].text:
                return True
    except: pass
    return False


# ─── Expériences significatives ───────────────────────────────────────────────

def _make_run(p_elem, text, bold=False, italic=False, color=None):
    r = etree.SubElement(p_elem, qn("w:r"))
    rpr = etree.SubElement(r, qn("w:rPr"))
    font = etree.SubElement(rpr, qn("w:rFonts"))
    font.set(qn("w:ascii"), "Calibri"); font.set(qn("w:hAnsi"), "Calibri")
    if bold:
        etree.SubElement(rpr, qn("w:b"))
        etree.SubElement(rpr, qn("w:bCs"))
    if italic:
        etree.SubElement(rpr, qn("w:i"))
        etree.SubElement(rpr, qn("w:iCs"))
    if color:
        c = etree.SubElement(rpr, qn("w:color")); c.set(qn("w:val"), color)
    sz = etree.SubElement(rpr, qn("w:sz")); sz.set(qn("w:val"), "24")
    szcs = etree.SubElement(rpr, qn("w:szCs")); szcs.set(qn("w:val"), "24")
    t = etree.SubElement(r, qn("w:t")); t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _build_exp_sig_paragraph(para, titre, societes, duree):
    p_elem = para._p
    for r in p_elem.findall(qn("w:r")):
        p_elem.remove(r)
    _make_run(p_elem, titre + " ", bold=True)
    _make_run(p_elem, "\u2013 ", bold=True, italic=True)
    _make_run(p_elem, societes.upper() + " ", bold=True, color="1F5C9A")
    _make_run(p_elem, "\u2013 ", bold=True, italic=True)
    _make_run(p_elem, duree.upper(), italic=True)


def _fill_experiences_significatives(tbl, exps):
    if not exps: return
    try:
        cell = tbl.rows[1].cells[0]
        paras = [p for p in cell.paragraphs if p.runs]
        for i, para in enumerate(paras):
            if i < len(exps):
                exp = exps[i]
                societes = exp.get("societes", exp.get("societe", ""))
                _build_exp_sig_paragraph(para, exp.get("titre",""), societes, exp.get("duree",""))
            else:
                for run in para.runs: run.text = ""
    except Exception as e:
        print(f"Erreur exp significatives: {e}")


# ─── Compétences ──────────────────────────────────────────────────────────────

def _fill_competences(tbl, data):
    try:
        cell = tbl.rows[1].cells[0]
        paras = list(cell.paragraphs)
        domaines   = data.get("competences_domaines", [])
        techniques = data.get("competences_techniques", [])
        print(f"[COMP] {len(paras)} paras, {len(domaines)} domaines, {len(techniques)} techniques")

        # Trouver tous les éléments clés AVANT toute écriture (évite faux positifs sur contenu réel)
        domain_title = next((p for p in paras if "Domaines" in get_text_in_element(p._p)), None)
        domain_start = paras.index(domain_title) + 1 if domain_title else 0
        domain_paras = paras[domain_start:domain_start + 5]

        tech_title = next((p for p in paras
                           if "techniques" in get_text_in_element(p._p).lower()), None)
        tech_title_idx = paras.index(tech_title) if tech_title else len(paras)
        tech_paras_found = paras[tech_title_idx + 1:tech_title_idx + 6] if tech_title_idx >= 0 else []

        print(f"[COMP] domain_title idx={paras.index(domain_title) if domain_title else None}")
        print(f"[COMP] domain_paras: {[get_text_in_element(p._p) for p in domain_paras]}")
        print(f"[COMP] tech_title idx={tech_title_idx}: {get_text_in_element(tech_title._p) if tech_title else None}")
        print(f"[COMP] tech_paras: {[get_text_in_element(p._p) for p in tech_paras_found]}")

        # ② Écrire "Domaine de Compétences" en gras
        if domain_title:
            _write_para_xml(domain_title, "Domaine de Compétences")
            runs_xml = list(domain_title._p.iter(qn("w:r")))
            if runs_xml:
                rPr = runs_xml[0].find(qn("w:rPr"))
                if rPr is None:
                    rPr = etree.SubElement(runs_xml[0], qn("w:rPr"))
                if rPr.find(qn("w:b")) is None:
                    etree.SubElement(rPr, qn("w:b"))

        # Écrire les 5 domaines avec police et interligne augmentés
        for i, para in enumerate(domain_paras):
            _write_para_xml(para, domaines[i] if i < len(domaines) else "")
            # Taille de police : 22 demi-points (11pt)
            for r in para._p.iter(qn("w:r")):
                rPr = r.find(qn("w:rPr"))
                if rPr is None:
                    rPr = etree.Element(qn("w:rPr"))
                    r.insert(0, rPr)
                for tag in (qn("w:sz"), qn("w:szCs")):
                    el = rPr.find(tag)
                    if el is None:
                        el = etree.SubElement(rPr, tag)
                    el.set(qn("w:val"), "22")
            # Interligne : 1.25x
            pPr = para._p.find(qn("w:pPr"))
            if pPr is None:
                pPr = etree.Element(qn("w:pPr"))
                para._p.insert(0, pPr)
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = etree.SubElement(pPr, qn("w:spacing"))
            spacing.set(qn("w:line"), "300")
            spacing.set(qn("w:lineRule"), "auto")

        # ③ Écrire "Compétences Techniques :"
        if tech_title:
            _write_para_xml(tech_title, "Compétences Techniques :")

        # Effacer les paras fantômes entre domaines et titre tech
        for p in paras[domain_start + 5:tech_title_idx]:
            _remove_bullet_style(p)

        if tech_title:
            _write_para_xml(tech_title, "Compétences Techniques :")

        # 5 lignes techniques
        tech_paras = tech_paras_found
        for i, para in enumerate(tech_paras):
            if i < len(techniques):
                cat = techniques[i].get("categorie", "")
                val = techniques[i].get("contenu", "")
                # Écrire "Catégorie :" dans les runs gras, valeur dans le run normal
                runs_xml = list(para._p.iter(qn("w:r")))
                bold_runs   = [r for r in runs_xml if r.find(qn("w:rPr")) is not None
                               and r.find(qn("w:rPr")).find(qn("w:b")) is not None]
                normal_runs = [r for r in runs_xml if r not in bold_runs]
                if bold_runs:
                    set_text_in_run(bold_runs[0], cat + " :")
                    for r in bold_runs[1:]:
                        for t in r.iter(qn("w:t")): t.text = ""
                elif runs_xml:
                    set_text_in_run(runs_xml[0], cat + " : " + val)
                if normal_runs:
                    set_text_in_run(normal_runs[0], " " + val)
                    for r in normal_runs[1:]:
                        for t in r.iter(qn("w:t")): t.text = ""
            else:
                for r in para._p.iter(qn("w:r")):
                    for t in r.iter(qn("w:t")): t.text = ""

    except Exception as e:
        traceback.print_exc()
        print(f"Erreur compétences: {e}")


# ─── Expériences professionnelles ─────────────────────────────────────────────

def _fill_experiences_pro(doc, exps):
    exp_tables = [t for t in doc.tables if _is_experience_pro_table(t)]
    if not exp_tables:
        return

    # Ajouter des tableaux manquants en copiant le dernier
    body = doc.element.body
    while len(exp_tables) < len(exps):
        last_tbl = exp_tables[-1]
        # Insérer un paragraphe vide + copie du tableau après le dernier tableau exp
        new_tbl = copy.deepcopy(last_tbl._tbl)
        sep = etree.Element(qn("w:p"))
        last_tbl._tbl.addnext(new_tbl)
        new_tbl.addprevious(sep)
        # Recharger la liste
        exp_tables = [t for t in doc.tables if _is_experience_pro_table(t)]

    for i, exp in enumerate(exps):
        _fill_single_exp_pro(exp_tables[i], exp)

    # Vider les tableaux en trop si le template en a plus que le CV
    for i in range(len(exps), len(exp_tables)):
        for row in exp_tables[i].rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs: run.text = ""


def _fill_single_exp_pro(tbl, exp):
    try:
        # Ligne 0 : dates | société
        row0 = tbl.rows[0]
        paras0 = row0.cells[0].paragraphs
        if len(paras0) >= 1 and paras0[0].runs:
            paras0[0].runs[0].text = exp.get("date_debut","") + " à"
            for r in paras0[0].runs[1:]: r.text = ""
        if len(paras0) >= 2 and paras0[1].runs:
            paras0[1].runs[0].text = exp.get("date_fin","")
            for r in paras0[1].runs[1:]: r.text = ""
        c1 = row0.cells[1]
        if c1.paragraphs and c1.paragraphs[0].runs:
            c1.paragraphs[0].runs[0].text = exp.get("societe","")
            for r in c1.paragraphs[0].runs[1:]: r.text = ""

        # Ligne 1 : poste | détail
        row1 = tbl.rows[1]
        for para in row1.cells[0].paragraphs:
            if para.runs:
                para.runs[0].text = exp.get("poste","")
                for r in para.runs[1:]: r.text = ""
                break
        _fill_exp_detail_cell(row1.cells[1], exp)
    except Exception as e:
        print(f"Erreur fill exp pro: {e}")


def _fill_exp_detail_cell(cell, exp):
    paras = cell.paragraphs

    def find_para(keyword):
        return next((p for p in paras if keyword.lower() in p.text.lower()), None)

    def write_para(para, text, bold=None):
        if para and para.runs:
            if bold is True:
                para.runs[0].bold = True
            elif bold is False:
                for r in para.runs: r.bold = False
            para.runs[0].text = text
            for r in para.runs[1:]: r.text = ""

    projet_para   = find_para("projet")
    missions_para = find_para("missions")
    env_para      = find_para("environnement")

    projet_idx   = paras.index(projet_para)   if projet_para   else -1
    missions_idx = paras.index(missions_para) if missions_para else -1
    env_idx      = paras.index(env_para)      if env_para      else len(paras)

    # "Projet :" en gras
    write_para(projet_para, "Projet :", bold=True)

    # Description du projet — paragraphe juste après "Projet :"
    if projet_idx >= 0 and projet_idx + 1 < len(paras):
        projets = exp.get("projets") or ([exp["projet"]] if exp.get("projet") else [])
        write_para(paras[projet_idx + 1], " / ".join(projets), bold=False)

    # "Missions :" en gras
    write_para(missions_para, "Missions :", bold=True)

    # Puces missions — travail en XML pur pour éviter les objets Paragraph invalidés
    missions = [m for m in exp.get("missions", []) if not _is_empty_or_placeholder(m)]

    if missions_idx >= 0 and missions_para and env_para:
        cell_xml = missions_para._p.getparent()
        all_p_xml = [c for c in cell_xml if c.tag == qn("w:p")]
        m_xml_idx   = all_p_xml.index(missions_para._p)
        env_xml_idx = all_p_xml.index(env_para._p)

        # Éléments bullet actuels entre "Missions :" et "Environnement"
        bullet_elems = all_p_xml[m_xml_idx + 1:env_xml_idx]

        # Capturer le pPr (puce/numPr) du premier bullet template pour l'uniformiser
        ref_ppr = None
        if bullet_elems:
            ppr = bullet_elems[0].find(qn("w:pPr"))
            if ppr is not None:
                ref_ppr = copy.deepcopy(ppr)

        # Capturer le rPr (police + taille) du paragraphe "Projet" comme référence
        ref_rpr = None
        if projet_para:
            for r in projet_para._p.iter(qn("w:r")):
                rpr = r.find(qn("w:rPr"))
                if rpr is not None:
                    ref_rpr = copy.deepcopy(rpr)
                    break

        # Ajouter des slots manquants (copie du premier bullet pour garder le style)
        while len(bullet_elems) < len(missions) and bullet_elems:
            new_p = copy.deepcopy(bullet_elems[0])
            env_para._p.addprevious(new_p)
            bullet_elems.append(new_p)

        # Supprimer les slots en trop
        for ep in bullet_elems[len(missions):]:
            cell_xml.remove(ep)
        bullet_elems = bullet_elems[:len(missions)]

        # Remplir chaque slot avec police/taille/puce uniformes
        for i, p_elem in enumerate(bullet_elems):
            # Normaliser le pPr (puce identique pour tous)
            if ref_ppr is not None:
                old_ppr = p_elem.find(qn("w:pPr"))
                if old_ppr is not None:
                    p_elem.remove(old_ppr)
                p_elem.insert(0, copy.deepcopy(ref_ppr))

            # Supprimer tous les runs existants
            for r in list(p_elem.findall(qn("w:r"))):
                p_elem.remove(r)

            # Créer un run avec le style du paragraphe "Projet" (sans gras)
            r = etree.SubElement(p_elem, qn("w:r"))
            if ref_rpr is not None:
                rpr = copy.deepcopy(ref_rpr)
                for bold_tag in (qn("w:b"), qn("w:bCs")):
                    b = rpr.find(bold_tag)
                    if b is not None:
                        rpr.remove(b)
                r.insert(0, rpr)
            t_elem = etree.SubElement(r, qn("w:t"))
            t_elem.text = missions[i]
            t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    # Saut de ligne entre dernière mission et environnement technique
    if env_para and bullet_elems:
        empty_sep = etree.Element(qn("w:p"))
        env_para._p.addprevious(empty_sep)

    # Environnement technique
    env = exp.get("environnement_technique", "")
    if env and env_para:
        normal_runs = [r for r in env_para.runs if not r.bold]
        if normal_runs:
            normal_runs[-1].text = " " + env

    # Max 1 saut de ligne après l'environnement technique — XML direct (évite le bug list.index)
    if env_para:
        cell_xml = env_para._p.getparent()
        all_p = [c for c in cell_xml if c.tag == qn("w:p")]
        env_xml_idx = all_p.index(env_para._p)
        for tp_xml in all_p[env_xml_idx + 2:]:
            cell_xml.remove(tp_xml)


# ─── Espacement global ────────────────────────────────────────────────────────

VML_TXBX = "{urn:schemas-microsoft-com:vml}textbox"
WPS_TXBX = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx"


def _find_txbx_para(body, keyword):
    """Trouve le paragraphe du body qui contient un rectangle avec le mot-clé donné."""
    for child in body:
        if child.tag == qn("w:p"):
            for elem in child.iter():
                if elem.tag in (VML_TXBX, WPS_TXBX):
                    text = "".join(t.text or "" for t in elem.iter(qn("w:t")))
                    if keyword.lower() in text.lower():
                        return child
    return None


def _set_page_break_before(para_elem):
    """Ajoute w:pageBreakBefore dans le pPr du paragraphe XML — Word force toujours un saut de page avant lui."""
    pPr = para_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.Element(qn("w:pPr"))
        para_elem.insert(0, pPr)
    if pPr.find(qn("w:pageBreakBefore")) is None:
        pPr.insert(0, etree.Element(qn("w:pageBreakBefore")))


def _remove_empty_paras_before(body, target_elem):
    """Supprime les paragraphes vides (sans texte ni shape) juste avant target_elem dans le body."""
    children = list(body)
    try:
        target_idx = children.index(target_elem)
    except ValueError:
        return
    for child in reversed(children[:target_idx]):
        if child.tag != qn("w:p"):
            break
        has_text = bool("".join(t.text or "" for t in child.iter(qn("w:t"))).strip())
        has_shape = any(e.tag in (VML_TXBX, WPS_TXBX) for e in child.iter())
        if has_text or has_shape:
            break
        body.remove(child)


def _remove_paras_before_first_exp(doc):
    """Force un saut de page avant le rectangle 'Expériences professionnelles'."""
    body = doc.element.body
    para = _find_txbx_para(body, "professionnelles")
    print(f"[EXP] para trouvé: {para is not None}")
    if para is not None:
        _remove_empty_paras_before(body, para)
        _set_page_break_before(para)


# ─── Espacement avant Formations ──────────────────────────────────────────────

def _fix_spacing_before_formations(doc):
    """Force un saut de page avant le rectangle 'Formations' et colle le tableau dessous."""
    body = doc.element.body

    formations_para = _find_txbx_para(body, "Formations")
    print(f"[FORM_SPACE] para trouvé: {formations_para is not None}")
    if formations_para is not None:
        _remove_empty_paras_before(body, formations_para)
        _set_page_break_before(formations_para)

    # Laisser exactement 1 paragraphe vide entre le rectangle et le tableau formations
    form_tbl = next((t for t in doc.tables if _is_formation_table(t)), None)
    if not form_tbl:
        return
    try:
        tbl_idx = list(body).index(form_tbl._tbl)
    except ValueError:
        return
    children = list(body)
    i = tbl_idx - 1
    while i >= 0 and children[i].tag == qn("w:p"):
        text = "".join(t.text or "" for t in children[i].iter(qn("w:t")))
        if not text.strip() and children[i] is not formations_para:
            body.remove(children[i])
        else:
            break
        i -= 1
    # Insérer exactement 1 paragraphe vide comme saut de ligne visuel
    tbl_idx_fresh = list(body).index(form_tbl._tbl)
    body.insert(tbl_idx_fresh, etree.Element(qn("w:p")))


# ─── Formations ───────────────────────────────────────────────────────────────

def _fill_formations(tbl, formations):
    # Ajouter des lignes si le CV en a plus que le template
    while len(tbl.rows) < len(formations):
        new_tr = copy.deepcopy(tbl.rows[-1]._tr)
        tbl._tbl.append(new_tr)

    rows = tbl.rows
    print(f"[FORM] {len(formations)} formations, {len(rows)} rows")
    for i, form in enumerate(formations):
        if i < len(rows):
            # Utiliser XML direct pour écrire dans les cellules (contourne les problèmes de para.runs)
            for cell, key in [(rows[i].cells[0], "annee"), (rows[i].cells[1], "intitule")]:
                for para in cell.paragraphs:
                    runs_xml = list(para._p.iter(qn("w:r")))
                    if runs_xml:
                        set_text_in_run(runs_xml[0], str(form.get(key) or ""))
                        for r in runs_xml[1:]:
                            for t in r.iter(qn("w:t")): t.text = ""
                        break

    for i in range(len(formations), len(tbl.rows)):
        for cell in tbl.rows[i].cells:
            for para in cell.paragraphs:
                for r in para._p.iter(qn("w:r")):
                    for t in r.iter(qn("w:t")): t.text = ""


# ─── Footer agence ────────────────────────────────────────────────────────────
# Le texte du footer est splitté entre plusieurs <w:t> dans le XML du docx.
# On inspecte le zip et on remplace les fragments exacts tels qu'ils apparaissent.

FOOTER_REPLACEMENTS = {
    # Agence Niort : DA = Béatrice HERITIER (06 62 01 25 58), reste inchangé
    # "Béatrice HERITIER" fait 4 chars de moins qu'Anne-Sophie MORANCAIS (17 vs 21)
    # → on ajoute 4 espaces de padding à la fin du run "HERITIER" pour réaligner
    "niort": [
        ("Anne-Sophie ",                       "B\u00e9atrice "),
        ("<w:t>MORANCAIS</w:t>",               "<w:t>HERITIER</w:t>"),
        # Le séparateur "- " est unique à la ligne DA (Robin/Pierre utilisent "-" sans espace)
        # On insère le padding ici — le preserve est déjà présent sur ce tag
        ('preserve">- </w:t>',                 'preserve">               - </w:t>'),  # +15 espaces avant le tiret
        ("06 82 30 40 25",                     "06 62 01 25 58"),
    ],
    # Agence Le Mans : adresse + IA1 = Benjamin BOUCHER (07 64 26 10 63)
    # "Benjamin BOUCHER" visuellement ~5 espaces plus large que Robin LAVOGEZ
    # → run 23 espaces → 15 espaces (total 18 vs 26 pour Robin, -8)
    "lemans": [
        ("Infotel Niort,",                     "Infotel Le Mans,"),
        (" 4 Bd Louis Tardy, 79000 NIORT",     " 2 Promenade d\u2019Androm\u00e8de, 72000 LE MANS"),
        ("Robin ",                             "Benjamin "),
        ("LAVOGEZ",                            "BOUCHER"),
        ("06 64 41 42 84",                     "07 64 26 10 63"),
        ("                       ",            "                   "),  # 23 → 19 espaces (total 22)
    ],
}


def _fill_footer_zip(docx_path: str, agence: str):
    """Patch footer*.xml directement dans le zip — contourne le problème de splits de w:t."""
    import io, zipfile as zf

    agence = agence.lower().strip()
    replacements = FOOTER_REPLACEMENTS.get(agence)
    if not replacements:
        print(f"[FOOTER] Agence '{agence}' : pas de remplacement défini.")
        return

    with open(docx_path, "rb") as fh:
        raw = fh.read()

    buf     = io.BytesIO(raw)
    out_buf = io.BytesIO()

    with zf.ZipFile(buf, "r") as zin, zf.ZipFile(out_buf, "w", zf.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if "footer" in item.filename.lower() and item.filename.endswith(".xml"):
                xml = data.decode("utf-8")
                for old, new in replacements:
                    n = xml.count(old)
                    if n:
                        xml = xml.replace(old, new)
                        print(f"[FOOTER] '{old}' → '{new}'  ({n}×)")
                data = xml.encode("utf-8")
            zout.writestr(item, data)

    with open(docx_path, "wb") as fh:
        fh.write(out_buf.getvalue())


# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx(data, agence: str = "niort"):
    # Copier le template dans un fichier temp pour ne jamais toucher l'original
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_template = tmp.name
    shutil.copy2(str(TEMPLATE_PATH), tmp_template)
    doc = Document(tmp_template)

    # ① Nettoyer "8 ANS" → "8" puis écrire "8 ans d'expérience"
    annees_raw   = str(data.get("annees_experience", ""))
    annees_clean = re.sub(r"\s*(ans|ANS|an|AN)\b.*", "", annees_raw).strip()
    replace_text_in_txbx(doc, "XX ans d\u2019exp\u00e9rience",
                          f"{annees_clean} ans d\u2019exp\u00e9rience")
    prenom = data.get("prenom", data.get("nom","")).upper()
    replace_text_in_txbx(doc, "CONSULTANT", prenom)
    replace_text_in_txbx(doc, "Nom du poste", data.get("poste",""))

    # Espacement EN PREMIER (les détecteurs cherchent les placeholders, avant tout remplissage)
    _remove_paras_before_first_exp(doc)
    _fix_spacing_before_formations(doc)

    for tbl in doc.tables:
        header = _get_table_header(tbl)
        if "Exp\u00e9riences significatives" in header:
            _fill_experiences_significatives(tbl, data.get("experiences_significatives",[]))
        elif "Comp\u00e9tences" in header or "Competences" in header:
            _fill_competences(tbl, data)
        elif _is_formation_table(tbl):
            _fill_formations(tbl, data.get("formations",[]))

    _fill_experiences_pro(doc, data.get("experiences_professionnelles",[]))

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        out_path = f.name

    # Patch footer directement dans le zip (après save) — le texte est splitté en w:t
    _fill_footer_zip(out_path, agence)

    return out_path


# ─── Routes API ───────────────────────────────────────────────────────────────

@app.post("/generate")
async def generate_cv(
    file: UploadFile = File(...),
    job_description: str = Form(default=""),
    influence: int = Form(default=0),
    agence: str = Form(default="niort"),
    user=Depends(get_current_user),
):
    content = await file.read()
    filename = file.filename.lower()

    if filename.endswith(".pdf"):
        text = extract_text_from_pdf(content)
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        text = extract_text_from_docx(content)
    else:
        raise HTTPException(400, "Format non supporté. Utilisez PDF ou DOCX.")

    if not text.strip():
        raise HTTPException(400, "Impossible d'extraire le texte du fichier.")

    try:
        data = extract_cv_data(text, job_description=job_description, influence=influence)
    except json.JSONDecodeError as e:
        traceback.print_exc()
        raise HTTPException(500, f"L'IA n'a pas retourné un JSON valide: {e}")
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(500, f"Erreur lors de l'appel IA: {e}")

    print(f"[AGENCE] agence={agence}")
    try:
        output_path = build_docx(data, agence=agence)
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(500, f"Erreur lors de la génération du DOCX: {e}")

    # ① Nom du fichier : FC PRENOM - Titre du poste - INFOTEL
    prenom     = data.get("prenom", "").strip().upper()
    titre      = data.get("poste", "Consultant").strip().title()
    fname      = f"FC {prenom} - {titre} - INFOTEL.docx" if prenom else f"FC - {titre} - INFOTEL.docx"

    log_generation(user["username"], fname)

    return FileResponse(output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=fname)


@app.get("/health")
def health():
    return {"status": "ok"}


@app.get("/", response_class=HTMLResponse)
def serve_frontend():
    return FRONTEND_PATH.read_text(encoding="utf-8")

@app.get("/admin-panel", response_class=HTMLResponse)
def serve_admin():
    return ADMIN_PATH.read_text(encoding="utf-8")


@app.get("/debug-competences")
def debug_competences():
    """Inspecte en détail les paragraphes de la cellule Compétences."""
    doc = Document(str(TEMPLATE_PATH))
    result = []
    for tbl in doc.tables:
        header = _get_table_header(tbl)
        if "Comp\u00e9tences" in header or "Competences" in header:
            cell = tbl.rows[1].cells[0]
            for i, p in enumerate(cell.paragraphs):
                result.append({
                    "idx": i,
                    "text": p.text,
                    "runs": [{"text": r.text, "bold": r.bold} for r in p.runs],
                    "xml_snippet": p._p.xml[:300],
                })
    return JSONResponse(result)


@app.get("/debug-template")
def debug_template():
    """Inspecte la structure réelle du template pour diagnostiquer les sélecteurs."""
    doc = Document(str(TEMPLATE_PATH))
    result = {"tables": [], "textboxes": []}

    for i, tbl in enumerate(doc.tables):
        tbl_info = {
            "index": i,
            "rows": len(tbl.rows),
            "cols": len(tbl.columns) if tbl.rows else 0,
            "preview": [],
        }
        for row in tbl.rows[:3]:
            tbl_info["preview"].append([cell.text[:120] for cell in row.cells])
        result["tables"].append(tbl_info)

    body = doc.element.body
    WPS = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx"
    VML = "{urn:schemas-microsoft-com:vml}textbox"
    for elem in body.iter():
        if elem.tag in (WPS, VML):
            texts = [get_text_in_element(p) for p in elem.iter(qn("w:p"))
                     if get_text_in_element(p).strip()]
            if texts:
                result["textboxes"].append(texts)

    return JSONResponse(result)
